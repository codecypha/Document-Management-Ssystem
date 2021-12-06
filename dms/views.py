from django.shortcuts import render, redirect
from dms.models import Department, Folder, FileType, VersionUpload, GroupManagement, AddUserGroup, AddUserDept, AddUserFolder, RelatedUpload
from .forms import DepartmentForm, FolderForm, FileForm, VersionForm, ManagementForm,AddUserForm, AddUserDeptForm, AddUserFolderForm, RelatedForm
from django.conf import settings
from django.core.exceptions import ObjectDoesNotExist
from django.core.exceptions import ObjectDoesNotExist
from django.contrib.auth.models import Group, User
# from pyad import *
# from pyad import aduser
from django.contrib.auth import login, logout, authenticate
from django.contrib.auth.decorators import login_required
from .decorators import unauthenticated_user, allowerd_users
from django.contrib import messages
from docx import Document
from django.contrib.staticfiles.utils import get_files
from django.contrib.staticfiles.storage import StaticFilesStorage
import os
from django.conf import settings
import win32com.client
import fitz
from django.db.models import Q
from itertools import chain
#from .tasks import job
from datetime import datetime
import datetime
from django.utils import timezone
from datetime import date
today = date.today()
d2  = today.strftime("%Y")

# Create your views here.
@unauthenticated_user
def loginPage(request): 
    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')
        user = authenticate(request, username=username, password=password)
        if user is not None:
            login(request, user)
            #messages.success(request, 'Incorrect Username Entered')
            messages.success(request, 'Welcome' + ' ' + username)
            return redirect('home') 
        else:
            messages.error(request, 'Incorrect Username Entered')
    context ={}
    return render(request, 'app/login.html', context)


@login_required(login_url  ='login')
def logoutPage(request):
    logout(request)
    return redirect('login')


@login_required(login_url  ='login')
def home(request):

    group_id = ''
    author = request.user  
    department_list = Department.objects.filter(author=author, status='active')
    department_count = Department.objects.filter(author=author, status='active').count()   
    folder_count = Folder.objects.filter(author=author).count()  
    file_count = FileType.objects.filter(author=author).count()    
    access_count = AddUserFolder.objects.raw(""" SELECT a.id, count(a.id) count FROM dms_adduserfolder a, dms_adduserdept b 
                                                    WHERE a.department_name = b.department_name
                                                    AND a.creator = b.creator """)
    for users in access_count:
        pass
    context = {"department_list":department_list, "group_id":group_id, "department_count":department_count, "users":users, "file_count":file_count, "folder_count":folder_count }
    return render(request, 'app/dashboard.html', context)



@login_required(login_url  ='login')
def department1(request):
    add_route ='department'
    group_id = ''
    my_user = request.user
    my_group = my_user.groups.all()
    for x in my_group:
        print(x)
    dept_users=[] 
    for dept_group in my_group:  
        dept_users.append(str(dept_group))
        dept_id = Department.objects.filter(group_name=dept_group, status='active').exists()
        if dept_id == True:
            group_id = 'allow_for_section'
         
            context = {"group_id":group_id} 
    department_count = Department.objects.filter(status='active').count()
    department_list = Department.objects.filter(status='active')
    form = DepartmentForm()
    form2 = AddUserForm()
    if request.method == 'POST' and 'adduser' not in request.POST:
        department_name = request.POST.get('department_name')
        #group_name = department_name + '_' + country
        group_name = department_name
        form = DepartmentForm(request.POST)
        #g=form.errors
        if form.errors:
            messages.error(request, "Section with this name " + department_name + " already exists")
        if form.is_valid():
            dept = form.save()
            dept.group_name = group_name
            dept.save()
            messages.success(request, "Section with the name " + department_name + " has been created")
            return redirect('department')

    if request.method == 'POST' and 'adduser' in request.POST:
        group_name = request.POST['adduser']
        username = request.POST.get('username')
        user = User.objects.get(username = username)
        user_group = Group.objects.get(name=group_name)
        user_group.user_set.add(user)
        messages.success(request, username + " Has been granted access.")
        return redirect('department')
    form = DepartmentForm()
         
    context = {"form":form, "form2":form2, "department_list":department_list, "department_count":department_count, "group_id":group_id, "add_route":add_route}
    return render(request, 'app/department.html', context)


@login_required(login_url  ='login')
def department(request):
    add_route ='department'
    groups_list=''
    request.session['add_route'] = add_route
    author  = request.user
    user_name = str(author)
    qs1 = Department.objects.filter(author=author, status='active')
    qs = AddUserDept.objects.filter(username=author) 
    for item in qs:
        department_name  = item.department_name
        owner = item.creator
        qs = AddUserDept.objects.filter(creator=owner, department_name=department_name, username=author)
    #qs = AddUserDept.objects.filter(creator=owner, department_name=department_name, username=author)
        if  qs.exists():
            groups_list = AddUserDept.objects.raw(""" SELECT distinct(a.id) id, a.department_name department_name, retention, retention_count, ACTION, 
                                                        period, a.entry_date, author FROM dms_department a, dms_adduserdept b, dms_adduserfolder c
                                                        WHERE a.department_name = b.department_name 
                                                        AND b.department_name = c.department_name
                                                        and b.username = %s
                                                        OR c.folder_access = 'partial' """, [user_name] )
        else:
            groups_list = []
      
    form = DepartmentForm()
    form2 = AddUserForm()
    status=''
    if request.method == 'POST' and 'adduser' not in request.POST:
        form = DepartmentForm(request.POST)
        group_name = request.POST.get('department_name')
        retention = request.POST.get('retention')   
        # if retention != 'Always Available':
        #     status = 'inactive'
        # else:
        #     status = 'active'
        request.session['group_name'] = group_name
        group_name = str(group_name)
        if form.is_valid():
 
            dept = form.save(commit=False)
            dept.group_name = group_name
            dept.author = author
           # dept.status = status
            dept.save()
            dept2 = AddUserDept(creator = author, username = author, department_name=group_name)
            dept2.save()
            return redirect('department')

    
    elif request.method == 'POST' and 'adduser' in request.POST:   
        groups_list = GroupManagement.objects.filter(creator=author)
        #group_name = request.session.get('group_name')
        username = request.POST.get('username')
        group_name = request.POST.get('groupname')
        user = User.objects.get(username = username)
        user_group = Group.objects.get(name=group_name)
        user_group.user_set.add(user)
        return redirect('department')
    context = {"form":form,"form2":form2, "groups_list":groups_list,  "add_route":add_route}
    return render(request, 'app/department.html', context)




@login_required(login_url  ='login')
def edit_department(request, pk):
    dept_id = Department.objects.get(id=pk)
    status= dept_id.retention
    form = DepartmentForm(instance=dept_id)
    if request.method == 'POST':
        form = DepartmentForm(request.POST, instance=dept_id)
        if form.is_valid():
            form.save()
            return redirect('department')
    context = {"form":form, "status":status}
    return render(request, 'app/edit-department.html', context)


@login_required(login_url  ='login')
def delete_department(request, pk):
    department_id = Department.objects.get(id=pk)
    if request.method == 'GET':
        department_id.delete()
        return redirect(department) 


@login_required(login_url  ='login')
def view_department1(request, pk):
    name = ''
    group_id = ''
    department_id=pk
    folder_count = Folder.objects.filter(department_id=pk).count()
    folder_list = Folder.objects.filter(department_id=pk)
    country = 'Nigeria'
    my_user = request.user
    my_group = my_user.groups.all()
    folder_users=[] 
    for folder_group in my_group:  
        folder_users.append(str(folder_group))
        dept_id = Folder.objects.filter(group_name=folder_group).exists()
        if dept_id == True:
            group_id = 'allow_for_folder'
            context = {"group_id":group_id}
    for doc in folder_list:
        dept_id = doc.department_id
        doc_log = Department.objects.get(id=dept_id)
        department_id = doc_log.id
        name = doc_log.department_name
        context = {"name":name}
    form = FolderForm()
    form2 = AddUserForm()
    if request.method == 'POST' and 'delete_folder' in request.POST and 'adduser' not in request.POST:
        print('oka1')
        folder_id=request.POST['delete_folder']
        #folder_id = folder_list.id
        folder = Folder.objects.get(id=folder_id)
        folder.delete()
        return redirect(view_department, pk=pk)
    if request.method == 'POST'and 'delete_folder' not in request.POST and 'adduser' not in request.POST:
        action = request.POST.get('action')
        department_name = request.POST.get('department_name')
        folder_name= request.POST.get('folder_name')
        group_name = folder_name + '_' + country
        form = FolderForm(request.POST)
        if form.is_valid():
            res = form.save()
            res.department_id = department_id
            res.group_name = group_name
            Group.objects.create(name=group_name)
            res.save()
            user = User.objects.get(username = my_user)
            user_group = Group.objects.get(name=group_name)
            user_group.user_set.add(user)
            messages.success(request, "Folder with the name " + folder_name + " has been created")
            form2 = AddUserForm()
            return redirect('view_department', pk=pk)
    if request.method == 'POST' and 'adduser' in request.POST and 'delete_folder' not in request.POST:
        group_name = request.POST['adduser']  
        username = request.POST.get('username')
        user = User.objects.get(username = username)
        user_group = Group.objects.get(name=group_name)
        user_group.user_set.add(user)
        messages.success(request, username + " Has been granted access.")
        return redirect('department')
            # return redirect('department')
    context = {"form":form, "folder_list":folder_list, "folder_count":folder_count, "name":name, "group_id":group_id, "form2":form2}
    return render(request, 'app/view-department.html', context)

@login_required(login_url  ='login')
def view_department(request, pk):
    department_id=pk
    add_route ='folder'
    request.session['add_route'] = add_route
    request.session['department_id'] = department_id
    author  = request.user
    user_name = str(author)
    folder_count = Folder.objects.filter(department_id=pk).count()  
    dept_qs = Department.objects.get(id=department_id)
    department_name  = dept_qs.department_name
    departmentname =  str(department_name)
    owner = dept_qs.author
    qs = AddUserDept.objects.filter(creator=owner, department_name=department_name, username=author)
    if  qs.exists():
        #folder_list = AddUserFolder.objects.filter(department_name=department_name, creator=owner )
        folder_list = AddUserFolder.objects.raw(""" SELECT distinct(a.id) id, a.folder_name folder_name, retention, retention_count, ACTION, 
                                period, a.entry_date, author FROM  dms_adduserfolder b, dms_folder a 
                                WHERE folder_access = 'full'
                                AND a.group_name = %s
                                AND a.folder_name = b.folder_name
                                and b.username= %s """, [departmentname, user_name] )

    else:
        folder_list = AddUserFolder.objects.raw(""" SELECT distinct(a.id) id, a.folder_name folder_name, retention, retention_count, ACTION, 
                                period, a.entry_date, author FROM  dms_adduserfolder b, dms_folder a 
                                WHERE folder_access = 'partial'
                                AND a.group_name = %s
                                AND a.folder_name = b.folder_name
                                and b.username= %s""", [departmentname, user_name] )
    form = FolderForm()
    form2 = AddUserForm()
    if request.method == 'POST' and 'adduser' not in request.POST:
        form = FolderForm(request.POST)
        folder_name = request.POST.get('folder_name')
        #group_name = request.session.get('group_name') 
        doc_id = Department.objects.get(id=department_id)
        group_name = str(doc_id.department_name)
        if form.is_valid():
            print(group_name)
            folder = form.save(commit=False)
            folder.group_name = group_name
            folder.department_id=department_id
            folder.author = author
            folder.save()
            folder = AddUserFolder(creator = author, username = author, folder_name=folder_name)
            folder.save()
            return redirect('department')
    elif request.method == 'POST' and 'adduser' in request.POST:   
        groups_list = GroupManagement.objects.filter(creator=author)
        #group_name = request.session.get('group_name')
        username = request.POST.get('username')
        group_name = request.POST.get('groupname')
        user = User.objects.get(username = username)
        user_group = Group.objects.get(name=group_name)
        user_group.user_set.add(user)
        return redirect('department')
    context = {"form":form,"form2":form2, "folder_list":folder_list, "folder_count":folder_count,  "add_route":add_route}
    return render(request, 'app/view-department.html', context)



val=None
@login_required(login_url  ='login')
def folder(request, pk):
    add_route ='files'
    request.session['add_route'] = add_route
    author = request.user
    fold_id = pk
    file_count = FileType.objects.filter(folder_id=pk).count()
    # try:
    file_list = FileType.objects.filter(folder_id=pk).order_by('entry_date')
    global val
    def val():
        return fold_id
    doc_log = Folder.objects.get(id=pk)
    department_id = doc_log.department_id
    folder_name = doc_log.folder_name
    folder_id = doc_log.id
    try:
        last_file =FileType.objects.latest('entry_date')
        document_no = last_file.document_id
        document_num = int(document_no)
        document_no = int(document_num + 1)
        fileid = FileType.objects.get(document_id=document_num)
        file_id = int(fileid.id + 1)
        file_id = fileid.id + 1
    except(ObjectDoesNotExist, TypeError)  as e:
        document_no = 1000
        file_id =  1
    form = FileForm(request.POST or None, request.FILES or None)
    if request.method == 'POST':
        tags = request.POST.get('tags')
        if form.is_valid():         
            res = form.save()
            res.document_id = document_no
            res.department_id = department_id
            res.folder_id = folder_id
            res.file_id = file_id
            res.tags = tags
            res.author = author
            res.save()
            #return redirect(folder, pk=fileid)
            form = FileForm()
    context = {"form":form, "file_list":file_list, "file_count":file_count,"folder_name":folder_name,}
    return render(request, 'app/folder.html', context)



@login_required(login_url  ='login')
def edit_folder(request, pk):
    
    folder_id = Folder.objects.get(id=pk)
    form = FolderForm(instance=folder_id)
    folder_log = Folder.objects.get(id=pk)
    department_id = folder_log.department_id
    status= folder_id.retention
    if request.method == 'POST':
        form = FolderForm(request.POST, instance=folder_id)
        if form.is_valid():
            res = form.save()
            res.department_id = department_id
            res.save()
            return redirect('department')
    context = {"form":form, 'status':status}
    return render(request, 'app/edit-folder.html', context)


@login_required(login_url  ='login')
def delete_folder(request, pk):
    folder_id = Folder.objects.get(id=pk)
    if request.method == 'GET':
        folder_id.delete()
        return redirect(department, pk=pk) 




@login_required(login_url  ='login')
def delete_file(request, pk):
    file_id = FileType.objects.get(id=pk)
    if request.method == 'GET':
        file_id.delete()
        return redirect(view-department, pk=pk) 


@login_required(login_url  ='login')
def view_file(request, pk):
    fileid = FileType.objects.get(id=pk)
    file_name = fileid.upload_file
    file_name = str(file_name)
    fileid1 = fileid.id
    document_id = fileid.document_id
    folder_id = fileid.folder_id
    file_id = fileid.file_id
    fileid_new = pk
    department_id = fileid.department_id
    versionid = VersionUpload.objects.filter(file_id=file_id)
    Relatedid = RelatedUpload.objects.filter(file_id=file_id)
    form = FileForm(instance=fileid)
    form2 = VersionForm(request.POST or None, request.FILES or None)
    if request.method == 'POST':
        form2 = VersionForm(request.POST or None, request.FILES or None)
        # filename = request.POST.get('new_upload')
        # filename = str(filename)
        # v1, v2 = file_name.split('.')
        # filename = v1+'.'+v2
        if form2.is_valid():         
            res = form2.save()
            res.document_id = document_id
            res.department_id = department_id
            res.folder_id = folder_id
            res.file_id = file_id
            #res.save()
            return redirect('department')
    context = {"form":form, "file_name":file_name,"document_id":document_id, 'fileid1':fileid1, "versionid":versionid, "fileid":fileid, "fileid_new":fileid_new, "Relatedid":Relatedid}
    return render(request, 'app/view-file.html', context)



@login_required(login_url  ='login')
def edit_file(request, pk):

    #fold_id = val()
    fileid = FileType.objects.get(id=pk)
    file_name = fileid.upload_file
    department_id = fileid.department_id
    folder_id = fileid.folder_id
    file_name = str(file_name)
    v1, v2 = file_name.split('.')
    upload_file2 = v1+'.'+v2
    fileid1 = fileid.id
    document_id = fileid.document_id
    folder_id = fileid.folder_id

    file_id = fileid.file_id
    versionid = VersionUpload.objects.filter(file_id=file_id)
    Relatedid = RelatedUpload.objects.filter(file_id=file_id)
    # for fileid2 in versionid:
    #     fileid2 = fileid.id
    #     print(fileid)
    form2 = VersionForm(request.POST or None, request.FILES or None)
    form3 = RelatedForm(request.POST or None, request.FILES or None)
    form = FileForm(instance=fileid)
    #department_id = folder_log.department_id
    if request.method == 'POST' and 'update' in request.POST:
        print("add update")
        v1 = request.POST.get('v1')
        v2 = '.'+ v2
        upload_file2 = v1 + v2
        form = FileForm(request.POST, instance=fileid)
        if form.is_valid():
            form = FileForm(request.POST, instance=fileid)
            res = form.save()
            res.department_id = department_id
            res.upload_file = upload_file2
            res.folder_id = folder_id
            res.save()
            return redirect(folder, pk=folder_id)
    if request.method == 'POST' and 'add_ver' in request.POST: 
        print("add version")
        if form2.is_valid():
            form2 = VersionForm(request.POST or None, request.FILES or None)
            res2 = form2.save()
            res2.document_id = document_id
            res2.department_id = department_id
            res2.folder_id = folder_id
            res2.file_id = file_id
            res2.file_name = file_name
            res2.save()
            return redirect(edit_file, pk=pk)
        
    if request.method == 'POST' and 'add_related' in request.POST: 
        print("add related")
        if form3.is_valid():
            form3 = RelatedForm(request.POST or None, request.FILES or None)
            res2 = form3.save()
            res2.document_id = document_id
            res2.department_id = department_id
            res2.folder_id = folder_id
            res2.file_id = file_id
            res2.file_name = file_name
            res2.save()
            return redirect(edit_file, pk=pk)
    context = {"form":form, "file_name":file_name, "v1":v1, "v2":v2, "document_id":document_id, 'fileid1':fileid1, "versionid":versionid, "Relatedid":Relatedid}
    return render(request, 'app/edit-file.html', context)



@login_required(login_url  ='login')
def upload_version(request):
    fileid = request.POST.get('fileid1')
    form2 = VersionForm(request.POST or None, request.FILES or None)
    if request.method == 'POST':
        fileid = request.POST.get('fileid1')
        form2 = VersionForm(request.POST or None, request.FILES or None)

        fileid = FileType.objects.get(id=fileid)
        file_name = fileid.upload_file
        file_id = fileid.file_id
        department_id = fileid.department_id
        folder_id = fileid.folder_id
        document_id = fileid.document_id
      
        if form2.is_valid():         
            res = form2.save()
            res.document_id = document_id
            res.department_id = department_id
            res.folder_id = folder_id
            res.file_id = file_id
            res.save()
            return redirect('view_file', id=fileid)
    context= {"form2":form2}
    return render(request, 'app/file-versions.html', context)





@login_required(login_url  ='login')
def CreateGroup(request):
    add_route ='group'
    request.session['add_route'] = add_route
    author  = request.user
    groups_list = GroupManagement.objects.filter(creator=author)
    form = ManagementForm()
    if request.method == 'POST' and 'adduser' not in request.POST:
        form = ManagementForm(request.POST)   
        group_name = request.POST.get('group_name')
        request.session['group_name'] = group_name
        user = User.objects.get(username = author)
        group_name = str(group_name)
        if form.is_valid():
            grp = form.save(commit=False)
            grp.creator = author
            grp.group_name = group_name
            grp.save()
            grp = AddUserGroup(creator = author, username = author, group_name=group_name)
            grp.save()
            # g1 = Group.objects.create(name=group_name)  
            # user_group = Group.objects.get(name=group_name)
            # user_group.user_set.add(user)
            return redirect('create')

    
    elif request.method == 'POST' and 'adduser' in request.POST:   
        groups_list = GroupManagement.objects.filter(creator=author)
        #group_name = request.session.get('group_name')
        username = request.POST.get('username')
        group_name = request.POST.get('groupname')
        user = User.objects.get(username = username)
        user_group = Group.objects.get(name=group_name)
        user_group.user_set.add(user)
  
        return redirect('create')
      
    # else:
    #     print('okay')
    #     # submit_value  = request.POST['get_group']
    #     # print(submit_value)
    #     return redirect('create')

    context = {"form":form, "groups_list":groups_list}
    return render(request, 'app/groupmanagement.html', context)


@login_required(login_url  ='login')
def add_to_group(request, pk):
    author  = request.user
    add_route = request.session.get('add_route') 
    department_id = request.session.get('department_id') 
   
    if add_route == 'group':
        list  = GroupManagement.objects.get(creator=author, id=pk)
        group_name = list.group_name
        group_list = AddUserGroup.objects.filter(creator=author, group_name=group_name)
    elif add_route == 'department':
        list  = Department.objects.get(author=author, id=pk) 
        group_name = list.department_name
        group_list = AddUserDept.objects.filter(creator=author, department_name=group_name)
    elif add_route == 'folder':
        list  = Folder.objects.get(author=author, id=pk)
        group_name = list.folder_name
        group_list = AddUserFolder.objects.filter(creator=author, folder_name=group_name)
    elif add_route == 'file':
        list  = FileType.objects.get(author=author, id=pk)
        group_name = list.file_name
        group_list = AddUserFolder.objects.filter(creator=author, file_name=group_name)
    else:
        pass
    #group_name = list.group_name

    form = AddUserForm
    if request.method == 'POST' :
        username = request.POST.get('username')
        if add_route == 'group':     
            grp = AddUserGroup(creator = author, username = username, group_name=group_name)
            grp.save()
        elif add_route == 'department':
            grp = AddUserDept(creator = author, username = username, department_name=group_name)
            grp.save()
        elif add_route == 'folder':
            dept_qs = Department.objects.get(id=department_id)
            department_name  = dept_qs.department_name
            qs = AddUserDept.objects.filter(creator=author, department_name=department_name, username=username)
            if  qs.exists():
                grp = AddUserFolder(creator = author, username = username, department_name=department_name, folder_name=group_name,)
                grp.save()
            else:
                grp = AddUserFolder(creator = author, username = username, department_name=department_name, folder_name=group_name, folder_access='partial' )
                grp.save()
        elif add_route == 'file':
            grp = AddUserDept(creator = author, username = username, department_name=group_name)
            {{ file.upload_file.url }}
            grp.save()    
        messages.success(request, username + " Has been granted access.")
        return redirect('department')
    context = {"form":form, "group_name": group_name, "group_list":group_list, "add_route":add_route}
    return render(request, 'app/adduser.html', context)


@login_required(login_url  ='login')
def searchdoc(request):
    filepath = 'media\Payment-Specification_update.pdf'
    with fitz.open(filepath) as doc:
        text = ""
        for page in doc:
            text += page.getText().strip()
            if 'VALUE_DATE6' in text:
                print('okay')
       # return text

    location2 = ''
    location1 = ''
    search_content1 = ""
    search_content2 = ""
    if request.method == 'POST':
        search = request.POST.get('search')
        search1 = search.capitalize()
        search2 = search.lower()
        search3 = search.upper() 
    arr = os.listdir('media')
    files = []
    for ar2 in arr:
        files.append(str(ar2))
    search_content1 =[]
    search_content3 =[]
    search_content4 = []
    search_content5= []
 
    for ar in arr:
        ar2=str(ar)
        if ".docx" in ar2 and "$" not in ar2 and ".doc" in ar2:
            filed = ar2
            document = Document("media/" + filed)
            location  = ("media/" + filed)
            location  = (filed)             
            for table in document.tables:
                filed = ar2
                document = Document("media/" + filed)               
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if search1 in paragraph.text or search2 in paragraph.text or search3 in paragraph.text:
                                location1  = location
                                if location1 not in search_content1:
                                    search_content1.append(str(location1))
                                # search_return1 = paragraph.text
                                # search_list1.append((search_return1))
                                # context ={"location1":location1}
                       
        
        elif ".pdf" in ar2 and "$" not in ar2:
            filed= ar2
            filepath  = ("media/" + filed)
            
            with fitz.open(filepath) as doc:
                text = ""
                for page in doc:
                    text += page.getText().strip()
                if search1 in text or search2 in text or search3 in text:
                    location1  = filed
                    if location1 not in search_content3:
                        search_content3.append(str(location1))

        elif ".jpg" in ar2 and "$" not in ar2 or ".png" in ar2 or ".bmp" in ar2 or ".gif" in ar2 or ".tiff" in ar2:
            try:
                from PIL import Image
            except ImportError:
                import Image
            import pytesseract
            filed= ar2
            filepath  = ("media/" + filed)
            pytesseract.pytesseract.tesseract_cmd = r'C:/Users/Kolawole Bayode/AppData/Local/Programs/Tesseract-OCR/tesseract.exe'
            text = pytesseract.image_to_string(Image.open(filepath))
            search_content4.append(filed)
        else:
            results = FileType.objects.filter(Q(tags__contains=search1) | Q(tags__contains=search2) | Q(tags__contains=search3) | Q(upload_file__contains=search1) | Q(upload_file__contains=search2) | Q(upload_file__contains=search3))
            results2  = FileType.objects.filter(Q(note__contains=search1) | Q( note__contains=search2) | Q( note__contains=search3) | Q( signed_by__contains=search1) | Q( signed_by__contains=search2) | Q( signed_by__contains=search3))
           # results3  = Folder.objects.filter(Q(folder_name__contains=search1) | Q( folder_name__contains=search2) | Q( folder_name__contains=search3))
            total_result = results | results2 
            for result in total_result:
                location = result.upload_file
                if location not in search_content5:
                    search_content5.append(location)
          


    search_content2 =[]
    for ar in arr:
        ar2=str(ar)
        if ".docx" in ar2 and "$" not in ar2 and ".doc" in ar2:
            filed = ar2
            document = Document("media/" + filed)
            #location  = ("media/" + filed)
            location  = (filed)
            for paragraph in document.paragraphs:
                if search1 in paragraph.text or search2 in paragraph.text or search3 in paragraph.text:
                    location2 = location
                    search_return2 = paragraph.text
                    if location2 not in search_content2:
                        search_content2.append(str(location2))
                    print(search_content2)
                    # search_return1 = paragraph.text
                    # search_list2.append((search_return1))
         
        

    search_content = search_content2 + search_content1  + search_content3 + search_content4 + search_content5 
    context ={"filed":filed, "ar2":ar2, "search_content":search_content}
    return render(request, 'app/searchdoc.html', context)

    
# import schedule
# import time

# def job():
#     print("I am working ...")
    
# schedule.every(10).seconds.do(job)
# schedule.every(10).minutes.do(job)
# schedule.every().hour.do(job)
# schedule.every().day.at("10:30").do(job)
# schedule.every(5).to(10).minutes.do(job)
# schedule.every().monday.do(job)
# schedule.every().monday.at("10:30").do(job)
# schedule.every().minute.at(":17").do(job)


# while 1:
#     schedule.run_pending()
#     time.sleep(1)